"""
Generate Word (.docx) documentation for 7 widgets.
Run: python3 generate_docs.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Helpers ────────────────────────────────────────────────────────────────

BRAND_BLUE  = RGBColor(0x4F, 0x6A, 0xFF)
DARK        = RGBColor(0x1E, 0x29, 0x3B)
GRAY        = RGBColor(0x47, 0x55, 0x69)
LIGHT_GRAY  = RGBColor(0xF1, 0xF5, 0xF9)
GREEN       = RGBColor(0x10, 0xB9, 0x81)
RED         = RGBColor(0xEF, 0x44, 0x44)

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.lstrip('#'))
    tcPr.append(shd)

def set_cell_borders(cell, color='E2E8F0'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def make_doc():
    doc = Document()
    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.font.color.rgb = DARK
    return doc

def add_title_block(doc, widget_name: str, subtitle: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(widget_name)
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = BRAND_BLUE

    p2 = doc.add_paragraph()
    r2 = p2.add_run(subtitle)
    r2.font.size = Pt(12)
    r2.font.color.rgb = GRAY
    doc.add_paragraph()

def add_h2(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = BRAND_BLUE

def add_h3(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = DARK

def add_body(doc, text: str):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(10)

def add_bullet(doc, text: str, bold_prefix: str = ''):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.bold = True
        rb.font.size = Pt(10)
    rt = p.add_run(text)
    rt.font.size = Pt(10)

def add_props_table(doc, rows: list):
    """rows = list of (prop, type, default, description) tuples"""
    headers = ['Параметр', 'Тип', 'Умолчание', 'Описание']
    widths  = [Cm(4.5), Cm(2.5), Cm(3.5), Cm(8.0)]

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = table.rows[0]
    for i, (cell, htext, w) in enumerate(zip(hdr.cells, headers, widths)):
        cell.width = w
        set_cell_bg(cell, '4F6AFF')
        set_cell_borders(cell, '3A56F0')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(htext)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for idx, (prop, typ, default, desc) in enumerate(rows):
        row = table.add_row()
        bg = 'FFFFFF' if idx % 2 == 0 else 'F8FAFC'
        values = [prop, typ, default, desc]
        for i, (cell, val) in enumerate(zip(row.cells, values)):
            cell.width = widths[i]
            set_cell_bg(cell, bg)
            set_cell_borders(cell)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if i == 0:
                run.font.color.rgb = RGBColor(0x2D, 0x3A, 0x8C)
                run.bold = True
            else:
                run.font.color.rgb = DARK

    doc.add_paragraph()

def add_note(doc, text: str, color: RGBColor = None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run('  ℹ  ' + text)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = color or GRAY


# ─── Widget docs ────────────────────────────────────────────────────────────

def doc_elem_button():
    doc = make_doc()
    add_title_block(doc, 'ElemButton', 'Виджет кнопки с навигацией, переключением состояний и эффектами')

    add_h2(doc, 'Назначение')
    add_body(doc, (
        'ElemButton — универсальный кликабельный элемент. Поддерживает переход '
        'по URL, генерацию событий платформы, копирование состояния страницы '
        'в буфер обмена, режим переключателя (toggle), состояние загрузки, '
        'условное отключение через переменную хранилища, а также богатый набор '
        'визуальных эффектов.'
    ))

    add_h2(doc, 'Основные возможности')
    add_bullet(doc, 'Переход по URL — внутренний, внешний, новая вкладка', '')
    add_bullet(doc, 'Генерация именованных событий платформы (`eventName`)', '')
    add_bullet(doc, 'Копирование URL страницы со всеми параметрами фильтров', '')
    add_bullet(doc, 'Копирование произвольного текста в буфер обмена', '')
    add_bullet(doc, 'Toggle-режим: кнопка переключает переменную хранилища между значениями', '')
    add_bullet(doc, 'Состояние загрузки (спиннер) на заданное время после клика', '')
    add_bullet(doc, 'Условное отключение через переменную хранилища (`btnDisableVar`)', '')
    add_bullet(doc, 'Hover-эффекты: default / lift / glow / scale / pulse / none', '')
    add_bullet(doc, 'Градиентный фон и стиль glassmorphism', '')
    add_bullet(doc, 'Иконки слева и справа от текста (MDI)', '')

    add_h2(doc, 'Параметры')

    add_h3(doc, 'Навигация и события')
    add_props_table(doc, [
        ('url',               'String',  '\'\'',     'URL перехода (поддерживает шаблонные переменные)'),
        ('eventName',         'Array',   '[]',        'Имена событий платформы, генерируемых при клике'),
        ('isTargetBlank',     'Boolean', 'false',     'Открывать URL в новой вкладке'),
        ('isClickSelf',       'Boolean', 'false',     'Переход на текущей странице (без перезагрузки)'),
        ('isCopyStore',       'Boolean', 'false',     'Копировать состояние экрана (все параметры) в URL'),
        ('isCopyStoreExcludeVars', 'String', '\'\'', 'Исключить переменные из скопированной ссылки (через запятую)'),
        ('shouldCopyText',    'Boolean', 'false',     'Копировать текст из поля `textToCopy` в буфер'),
        ('textToCopy',        'String',  '\'\'',     'Текст для копирования в буфер обмена'),
        ('cutParams',         'Array',   '[]',        'Параметры URL, которые нужно убрать при переходе'),
        ('filters',           'Array',   '[]',        'Фильтры, передаваемые при переходе'),
        ('hasClassName',      'String',  '\'\'',     'CSS-класс, применяемый к элементу :has() при клике'),
    ])

    add_h3(doc, 'Текст и иконки')
    add_props_table(doc, [
        ('btnShowText',     'Boolean', 'true',        'Показывать подпись кнопки'),
        ('btnText',         'String',  '\'Подробнее\'','Текст кнопки'),
        ('btnIconLeft',     'String',  '\'\'',        'MDI-иконка слева от текста (имя без префикса mdi-)'),
        ('btnIconRight',    'String',  '\'\'',        'MDI-иконка справа от текста'),
    ])

    add_h3(doc, 'Внешний вид')
    add_props_table(doc, [
        ('btnBg',             'String',  '#4f6aff',   'Фоновый цвет кнопки'),
        ('btnColor',          'String',  '#ffffff',   'Цвет текста'),
        ('btnBorderRadius',   'String',  '8px',       'Скругление углов'),
        ('btnFontSize',       'String',  '14px',      'Размер шрифта'),
        ('btnFontWeight',     'String',  '500',       'Толщина шрифта'),
        ('btnFontFamily',     'String',  '\'\'',      'Гарнитура шрифта'),
        ('btnLetterSpacing',  'String',  '\'\'',      'Межбуквенный интервал'),
        ('btnTextTransform',  'String',  'none',      'Трансформация текста (uppercase / lowercase / none)'),
        ('btnPaddingV',       'String',  '10px',      'Вертикальный отступ'),
        ('btnPaddingH',       'String',  '20px',      'Горизонтальный отступ'),
        ('btnShadow',         'String',  '...',       'CSS box-shadow кнопки'),
        ('btnBorderWidth',    'String',  '0px',       'Толщина рамки'),
        ('btnBorderColor',    'String',  'transparent','Цвет рамки'),
        ('btnCustomCss',      'String',  '\'\'',      'Произвольный CSS для кнопки'),
        ('btnHoverCss',       'String',  '\'\'',      'Произвольный CSS при наведении'),
        ('btnGradientTo',     'String',  '\'\'',      'Конечный цвет градиента (если задан — активируется градиент)'),
        ('btnGradientAngle',  'String',  '135deg',    'Угол градиента'),
        ('btnIsGlass',        'Boolean', 'false',     'Стиль glassmorphism (размытый прозрачный фон)'),
        ('btnHoverEffect',    'String',  'default',   'Hover-эффект: default / lift / glow / scale / pulse / none'),
        ('btnCursor',         'String',  'pointer',   'CSS cursor при наведении'),
    ])

    add_h3(doc, 'Toggle-режим')
    add_props_table(doc, [
        ('btnIsToggle',          'Boolean', 'false', 'Включить режим переключателя'),
        ('btnToggleStoreVar',    'String',  '\'\'',  'Имя переменной хранилища для хранения состояния'),
        ('btnToggleActiveValue', 'String',  '\'\'',  'Значение переменной, при котором кнопка «активна»'),
        ('btnToggleBg',          'String',  '#1e293b','Фон в активном состоянии'),
        ('btnToggleColor',       'String',  '#ffffff','Цвет текста в активном состоянии'),
    ])

    add_h3(doc, 'Загрузка и отключение')
    add_props_table(doc, [
        ('btnLoadingOnClick',   'Boolean', 'false', 'Показывать спиннер загрузки после клика'),
        ('btnLoadingDuration',  'Number',  '1500',  'Длительность состояния загрузки (мс)'),
        ('btnDisableVar',       'String',  '\'\'',  'Переменная хранилища: если truthy — кнопка недоступна'),
    ])

    doc.save('ElemButton.docx')
    print('✓ ElemButton.docx')


def doc_elem_calendar():
    doc = make_doc()
    add_title_block(doc, 'ElemCalendar', 'Виджет календаря с выбором дат, событиями и тепловой картой')

    add_h2(doc, 'Назначение')
    add_body(doc, (
        'ElemCalendar — многофункциональный календарный виджет с двумя основными режимами: '
        'полный (full) — полноразмерный календарь с видами месяц/неделя/день/повестка/год, '
        'и компактный (compact) — встраиваемый фильтр дат. '
        'Поддерживает выбор одиночной даты, диапазона и множественного выбора, '
        'отображение событий из источника данных, тепловую карту метрик.'
    ))

    add_h2(doc, 'Режимы')
    add_h3(doc, 'Полный режим (calMode = "full")')
    add_bullet(doc, 'Переключение видов: месяц, неделя, день, повестка, год', '')
    add_bullet(doc, 'Отображение событий из хранилища или JSON', '')
    add_bullet(doc, 'Тепловая карта с настраиваемыми цветами', '')
    add_bullet(doc, 'Тултип при наведении на день с событиями', '')
    add_bullet(doc, 'Кнопки «Сегодня» и «Сброс»', '')

    add_h3(doc, 'Компактный режим (calMode = "compact") — 4 подрежима')
    add_props_table(doc, [
        ('Выбор периода',              '(стандартный)', 'Мини-календарь + пресеты', 'Базовый вид фильтра дат'),
        ('Выбор периода (расширенный)', '(extended)',   'Два мини-календаря + пресеты', 'Независимая навигация каждого месяца'),
        ('Выбор периода (упрощённый)', '(simplified)', 'Только пресеты (без календаря)', 'Быстрый выбор через кнопки-пресеты'),
        ('Ввод дат',                   '(input)',       'Только поля ввода дат', 'Ручной ввод начала и конца диапазона'),
    ])

    add_h2(doc, 'Параметры')

    add_h3(doc, 'Основные')
    add_props_table(doc, [
        ('calMode',           'String',  'full',       'Режим виджета: full | compact'),
        ('calView',           'String',  'month',      'Начальный вид (month / week / day / agenda / year)'),
        ('calLocale',         'String',  'ru',         'Язык интерфейса'),
        ('calFirstDay',       'Number',  '1',          'Первый день недели: 0=Вс, 1=Пн'),
        ('calShowWeekNumbers','Boolean', 'false',      'Показывать номера недель'),
        ('calShowWeekends',   'Boolean', 'true',       'Показывать выходные дни'),
        ('calShowHeader',     'Boolean', 'true',       'Показывать заголовок'),
        ('calShowViewSwitcher','Boolean','true',       'Показывать переключатель видов'),
        ('calShowTodayBtn',   'Boolean', 'true',       'Показывать кнопку «Сегодня»'),
        ('calShowResetBtn',   'Boolean', 'true',       'Полный режим: показывать кнопку «Сброс»'),
        ('calAvailableViews', 'Array',   '...',        'Список доступных видов'),
    ])

    add_h3(doc, 'Выбор дат')
    add_props_table(doc, [
        ('calSelectionMode',  'String', 'range',   'Режим выбора: none / single / range / multi'),
        ('calSelectedDate',   'String', '\'\'',    'Выбранная дата (ISO, режим single)'),
        ('calSelectedStart',  'String', '\'\'',    'Начало диапазона (ISO)'),
        ('calSelectedEnd',    'String', '\'\'',    'Конец диапазона (ISO)'),
        ('calSelectedDates',  'String', '\'\'',    'Выбранные даты (JSON-массив ISO, режим multi)'),
        ('calWithTime',       'Boolean','false',   'Режим с временем: передаёт timestamp в мс'),
        ('calDefaultStartTime','String','00:00',   'Время начала по умолчанию (HH:MM)'),
        ('calDefaultEndTime', 'String', '23:59',   'Время конца по умолчанию (HH:MM)'),
    ])

    add_h3(doc, 'Компактный режим')
    add_props_table(doc, [
        ('calCompactShowBottom',   'Boolean', 'true',  'Показывать нижнюю панель (поля ввода + кнопки)'),
        ('calCompactShowPresets',  'Boolean', 'true',  'Показывать кнопки-пресеты'),
        ('calCompactShowCalendar', 'Boolean', 'true',  'Показывать мини-календарь'),
        ('calCompactDualMonth',    'Boolean', 'false', 'Показывать два месяца рядом (расширенный режим)'),
        ('calCompactShowToday',    'Boolean', 'true',  'Показывать кнопку «Сегодня» в шапке (отключается в расширенном режиме)'),
        ('calCompactShowTooltip',  'Boolean', 'false', 'Тултип при наведении на день в компактном режиме'),
        ('calPresetsJson',         'String',  '...',   'JSON-массив пресетов [{key, label}]. Ключи: today, yesterday, week, last_week, month, last_month'),
        ('calPresetsColumns',      'Number',  '3',     'Число колонок пресетов (0 = авто)'),
        ('calCompactChipsGap',     'String',  '4px',   'Отступ между чипами «С» / «По»'),
    ])

    add_note(doc, (
        'В упрощённом режиме пресеты с ключами today и yesterday автоматически скрываются. '
        'Если добавить кастомный пресет с таким ключом — он тоже не отобразится в упрощённом режиме.'
    ))

    add_h3(doc, 'События и данные')
    add_props_table(doc, [
        ('calEventsVar',        'String', '\'\'', 'Переменная хранилища с событиями (JSON)'),
        ('calEventsJson',       'String', '\'\'', 'Статические события (JSON)'),
        ('calDataTitleCol',     'String', '\'\'', 'Столбец заголовка события в датасете'),
        ('calDataStartTimeCol', 'String', '\'\'', 'Столбец времени начала'),
        ('calDataEndTimeCol',   'String', '\'\'', 'Столбец времени конца'),
        ('calDataDescCol',      'String', '\'\'', 'Столбец описания события'),
    ])

    add_h3(doc, 'Тепловая карта')
    add_props_table(doc, [
        ('calHeatmapEnabled',   'Boolean', 'false',   'Включить тепловую карту'),
        ('calMetricVar',        'String',  '\'\'',    'Переменная хранилища с метрикой'),
        ('calMetricJson',       'String',  '\'\'',    'Статическая метрика (JSON)'),
        ('calMetricDataCol',    'String',  '\'\'',    'Столбец метрики в датасете'),
        ('calHeatmapColorLow',  'String',  '#f0f9ff', 'Цвет минимума'),
        ('calHeatmapColorHigh', 'String',  '#4f46e5', 'Цвет максимума'),
        ('calHeatmapShowValue', 'Boolean', 'true',    'Показывать значение в ячейке'),
    ])

    add_h3(doc, 'Оформление')
    add_props_table(doc, [
        ('calPreset',          'String', 'corporate', 'Пресет оформления'),
        ('calBg',              'String', '#ffffff',   'Фон виджета'),
        ('calHeaderBg',        'String', '#4f6aff',   'Фон шапки'),
        ('calHeaderColor',     'String', '#ffffff',   'Цвет текста шапки'),
        ('calAccentColor',     'String', '#4f6aff',   'Акцентный цвет'),
        ('calTodayBg',         'String', '#eef0ff',   'Фон ячейки «Сегодня»'),
        ('calTodayColor',      'String', '#4f6aff',   'Цвет текста «Сегодня»'),
        ('calWeekendColor',    'String', '#ef4444',   'Цвет выходных дней'),
        ('calSelectedBg',      'String', '#4f6aff',   'Фон выбранной даты'),
        ('calSelectedColor',   'String', '#ffffff',   'Цвет текста выбранной даты'),
        ('calRangeBg',         'String', 'rgba(79,106,255,0.10)', 'Фон диапазона'),
        ('calCellBg',          'String', '#ffffff',   'Фон ячейки'),
        ('calCellHoverBg',     'String', '#f5f7ff',   'Фон ячейки при наведении'),
        ('calCellBorderColor', 'String', '#e8ecf4',   'Цвет границы ячейки'),
        ('calRadius',          'String', '12px',      'Скругление виджета'),
        ('calEventRadius',     'String', '4px',       'Скругление метки события'),
        ('calShadow',          'String', '...',       'Тень виджета (CSS box-shadow)'),
        ('calFontFamily',      'String', 'inherit',   'Гарнитура шрифта'),
        ('calFontSize',        'String', '13px',      'Размер шрифта'),
        ('calTooltipBg',       'String', '#1e293b',   'Фон тултипа'),
        ('calTooltipColor',    'String', '#f1f5f9',   'Цвет текста тултипа'),
        ('calTooltipRadius',   'String', '10px',      'Скругление тултипа'),
        ('calCustomCss',       'String', '\'\'',      'Произвольный CSS'),
    ])

    add_h2(doc, 'Переменные хранилища (выходные)')
    add_props_table(doc, [
        ('date',      '—', '—', 'Выбранная дата (ISO-строка или timestamp мс при calWithTime=true)'),
        ('dateStart', '—', '—', 'Начало диапазона (ISO или timestamp)'),
        ('dateEnd',   '—', '—', 'Конец диапазона (ISO или timestamp)'),
        ('datesList', '—', '—', 'Массив выбранных дат в JSON-строке (режим multi)'),
    ])
    add_note(doc, 'Переменные настраиваются через системный раздел «Переменные» виджета на платформе.')

    doc.save('ElemCalendar.docx')
    print('✓ ElemCalendar.docx')


def doc_elem_sticker():
    doc = make_doc()
    add_title_block(doc, 'ElemSticker', 'Виджет стикера — интерактивная заметка с богатым оформлением')

    add_h2(doc, 'Назначение')
    add_body(doc, (
        'ElemSticker имитирует бумажный стикер (Post-it). Поддерживает inline WYSIWYG-редактирование, '
        '15 цветовых тем, булавку и ленту как декоративные элементы, текстуры бумаги, '
        'рукописные шрифты, метаданные (автор, приоритет, теги, даты) и эффекты Figma '
        '(fills / strokes / effects).'
    ))

    add_h2(doc, 'Основные возможности')
    add_bullet(doc, '15 готовых тем: yellow, pink, blue, green, purple, orange, red, teal, mint, lavender, peach, sky, rose, amber, lime', '')
    add_bullet(doc, 'Декоративные элементы: булавка (4 позиции) и скотч (4 позиции, 3 стиля)', '')
    add_bullet(doc, 'Текстуры бумаги: standard / aged / recycled + паттерны: lined / grid / dotted', '')
    add_bullet(doc, 'Загнутый уголок (4 позиции) и тень по краю', '')
    add_bullet(doc, 'Рукописные шрифты: Caveat, Kalam, Patrick Hand, Dancing Script, Indie Flower', '')
    add_bullet(doc, 'Метаданные: автор, дата создания/изменения, приоритет, теги', '')
    add_bullet(doc, 'Произвольный поворот (−10°..+10°) + случайный поворот при монтировании', '')

    add_h2(doc, 'Параметры')

    add_h3(doc, 'Контент')
    add_props_table(doc, [
        ('html',            'String',  '\'\'',          'HTML-контент стикера (редактируется inline)'),
        ('placeholder',     'String',  '\'Нажмите...\'','Подсказка при пустом контенте'),
        ('isPlainHTMLShown','Boolean', 'false',          'Показывать исходный HTML вместо WYSIWYG'),
    ])

    add_h3(doc, 'Тема и цвета')
    add_props_table(doc, [
        ('stickerTheme',  'String',  'yellow', 'Предустановленная цветовая тема (15 вариантов)'),
        ('stickerColor',  'String',  '\'\'',   'Кастомный фоновый цвет (переопределяет тему)'),
        ('textColor',     'String',  '\'\'',   'Кастомный цвет текста (переопределяет тему)'),
        ('useGradient',   'Boolean', 'true',   'Использовать градиент для объёмности фона'),
        ('showBorder',    'Boolean', 'false',  'Тонкая рамка стикера'),
    ])

    add_h3(doc, 'Поворот')
    add_props_table(doc, [
        ('rotation',       'Number',  '0',     'Угол поворота (от −10 до 10 градусов)'),
        ('rotationRandom', 'Boolean', 'false', 'Случайный небольшой поворот при монтировании'),
    ])

    add_h3(doc, 'Булавка')
    add_props_table(doc, [
        ('showPin',     'Boolean', 'true',        'Показывать декоративную булавку'),
        ('pinPosition', 'String',  'top-center',  'Позиция: top-left / top-center / top-right / center'),
        ('pinColor',    'String',  '\'\'',        'Цвет булавки (по умолчанию из темы)'),
    ])

    add_h3(doc, 'Скотч (tape)')
    add_props_table(doc, [
        ('showTape',     'Boolean', 'false',      'Показывать декоративный скотч'),
        ('tapeStyle',    'String',  'single',     'Стиль: single / double / torn'),
        ('tapePosition', 'String',  'top-center', 'Позиция: top-left / top-center / top-right / center'),
        ('tapeColor',    'String',  '\'\'',       'Цвет скотча'),
        ('tapeRotation', 'Number',  '0',          'Поворот скотча (от −15 до 15 градусов)'),
    ])

    add_h3(doc, 'Стиль бумаги')
    add_props_table(doc, [
        ('paperStyle',         'String',  'standard',     'Фон бумаги: standard / aged / recycled'),
        ('paperTexture',       'String',  'none',         'Текстура: none / lined / grid / dotted'),
        ('showFoldedCorner',   'Boolean', 'true',         'Загнутый уголок'),
        ('foldedCornerPosition','String', 'bottom-right', 'Позиция уголка: bottom-right / bottom-left / top-right / top-left'),
        ('edgeShadow',         'Boolean', 'false',        'Тень приподнятого края'),
    ])

    add_h3(doc, 'Шрифт')
    add_props_table(doc, [
        ('useHandwritingFont', 'Boolean', 'false',   'Рукописный шрифт'),
        ('handwritingFont',    'String',  'caveat',  'Гарнитура: caveat / kalam / patrick-hand / dancing-script / indie-flower'),
        ('fontFamily',         'String',  '\'\'',    'Произвольная гарнитура (переопределяет всё)'),
        ('fontSize',           'String',  '14px',    'Размер шрифта'),
        ('fontWeight',         'Number',  '400',     'Толщина шрифта'),
        ('lineHeight',         'Number',  '1.5',     'Межстрочный интервал'),
        ('textAlign',          'String',  'left',    'Выравнивание текста'),
    ])

    add_h3(doc, 'Метаданные')
    add_props_table(doc, [
        ('showMetadata',    'Boolean', 'true',   'Показывать блок метаданных'),
        ('metadataPosition','String',  'bottom', 'Позиция: top / bottom'),
        ('createdAt',       'String',  '\'\'',   'Дата создания (ISO)'),
        ('modifiedAt',      'String',  '\'\'',   'Дата изменения (ISO)'),
        ('author',          'String',  '\'\'',   'Автор заметки'),
        ('dateFormat',      'String',  'DD MMM YYYY', 'Формат даты (dayjs)'),
        ('showPriority',    'Boolean', 'false',  'Показывать индикатор приоритета'),
        ('priority',        'String',  '\'\'',   'Уровень приоритета'),
        ('showTags',        'Boolean', 'false',  'Показывать теги'),
        ('tags',            'Array',   '[]',     'Массив строк-тегов'),
    ])

    add_h3(doc, 'Размеры и отступы')
    add_props_table(doc, [
        ('padding',     'String', '16px', 'Внутренний отступ'),
        ('minHeight',   'String', '60px', 'Минимальная высота'),
        ('minWidth',    'String', '\'\'', 'Минимальная ширина'),
        ('maxWidth',    'String', '\'\'', 'Максимальная ширина'),
        ('aspectRatio', 'String', '\'\'', 'Соотношение сторон (CSS aspect-ratio)'),
        ('borderRadius','String', '2px',  'Скругление углов'),
        ('boxShadow',   'String', '...',  'Тень (CSS box-shadow)'),
        ('opacity',     'Number', '1',    'Прозрачность (0–1)'),
    ])

    doc.save('ElemSticker.docx')
    print('✓ ElemSticker.docx')


def doc_elem_text():
    doc = make_doc()
    add_title_block(doc, 'ElemText', 'Виджет текста с типографикой, Figma-совместимыми эффектами и шаблонными переменными')

    add_h2(doc, 'Назначение')
    add_body(doc, (
        'ElemText — полнофункциональный текстовый блок с inline WYSIWYG-редактированием. '
        'Поддерживает дизайн-систему Insight, полный набор типографических свойств, '
        'Figma-совместимые fills / strokes / effects, шаблонные переменные {{var}}, '
        'усечение строк и вертикальное выравнивание.'
    ))

    add_h2(doc, 'Основные возможности')
    add_bullet(doc, 'Inline-редактирование HTML-контента прямо в конструкторе', '')
    add_bullet(doc, 'Шаблонные переменные: {{имя_переменной}} подставляются из хранилища', '')
    add_bullet(doc, 'Усечение текста: truncate=true (одна строка + «…»), truncate=N (N строк)', '')
    add_bullet(doc, 'Вертикальное выравнивание: top / center / bottom', '')
    add_bullet(doc, 'Figma-совместимые fills, strokes и effects (тени, размытие)', '')
    add_bullet(doc, 'Предустановки дизайн-системы Insight', '')
    add_bullet(doc, 'Произвольный HTML-тег корневого элемента (div, span, p, h1–h6...)', '')

    add_h2(doc, 'Параметры')

    add_h3(doc, 'Контент')
    add_props_table(doc, [
        ('html',            'String',  '\'Текст...\'', 'HTML-контент (поддерживает {{переменные}})'),
        ('tag',             'String',  'div',          'HTML-тег корневого элемента'),
        ('placeholder',     'String',  '\'Click to edit...\'', 'Подсказка в редакторе'),
        ('isPlainHTMLShown','Boolean', 'false',        'Режим исходного HTML'),
        ('editable',        'Boolean', 'true',         'Разрешить редактирование в конструкторе'),
        ('selectable',      'Boolean', 'true',         'Разрешить выделение текста'),
    ])

    add_h3(doc, 'Типографика')
    add_props_table(doc, [
        ('designSystem',    'String',  'insight',  'Предустановка дизайн-системы'),
        ('textAlign',       'String',  'left',     'Горизонтальное выравнивание: left / center / right / justify'),
        ('verticalAlign',   'String',  'top',      'Вертикальное выравнивание: top / center / bottom'),
        ('fontSize',        'String',  '\'\'',     'Размер шрифта (CSS-значение)'),
        ('fontWeight',      'String',  '\'\'',     'Толщина шрифта'),
        ('fontFamily',      'String',  '\'\'',     'Гарнитура шрифта'),
        ('lineHeight',      'String',  '\'\'',     'Межстрочный интервал'),
        ('letterSpacing',   'String',  '\'\'',     'Межбуквенный интервал'),
        ('textTransform',   'String',  '\'\'',     'Трансформация: none / uppercase / lowercase / capitalize'),
        ('textShadow',      'String',  '\'\'',     'Тень текста (CSS text-shadow)'),
        ('textDecoration',  'String',  '\'\'',     'Подчёркивание / зачёркивание'),
        ('truncate',        'Boolean/Number','false','true = одна строка + «…», N = N строк с «…»'),
        ('size',            'String',  'medium',   'Размер из дизайн-системы: xs / sm / medium / lg / xl'),
        ('weight',          'String',  'normal',   'Вес из дизайн-системы: light / normal / medium / semibold / bold'),
    ])

    add_h3(doc, 'Цвета и фон')
    add_props_table(doc, [
        ('color',           'String', '\'\'', 'Цвет текста (CSS)'),
        ('backgroundColor', 'String', '\'\'', 'Фоновый цвет блока'),
        ('linkColor',       'String', '\'\'', 'Цвет ссылок'),
        ('linkHoverColor',  'String', '\'\'', 'Цвет ссылок при наведении'),
    ])

    add_h3(doc, 'Отступы и рамка')
    add_props_table(doc, [
        ('padding',       'String', '\'\'',   'Внутренний отступ (CSS)'),
        ('margin',        'String', '\'\'',   'Внешний отступ (CSS)'),
        ('borderColor',   'String', '\'\'',   'Цвет рамки'),
        ('borderWidth',   'String', '\'\'',   'Толщина рамки (CSS)'),
        ('borderStyle',   'String', '\'\'',   'Стиль рамки (solid / dashed / dotted)'),
        ('borderRadius',  'String', '\'\'',   'Скругление (CSS)'),
        ('boxShadow',     'String', '\'\'',   'Тень (CSS box-shadow)'),
        ('opacity',       'Number', '\'\'',   'Прозрачность (0–1)'),
    ])

    add_h3(doc, 'Figma-совместимость')
    add_props_table(doc, [
        ('fills',   'Array', '[]', 'Слои заливки (Figma fills): [{type, color, opacity, ...}]'),
        ('strokes', 'Array', '[]', 'Слои обводки (Figma strokes)'),
        ('effects', 'Array', '[]', 'Эффекты: тени DROP_SHADOW / INNER_SHADOW, размытие LAYER_BLUR'),
        ('textAutoResize', 'String', 'NONE', 'Авторазмер текста (Figma): NONE / WIDTH_AND_HEIGHT / HEIGHT'),
    ])

    add_h3(doc, 'Переменные хранилища')
    add_props_table(doc, [
        ('varBindings', 'Object', '{}', 'Привязка переменных: {listen: {}, write: {}} — двусторонняя синхронизация'),
    ])
    add_note(doc, 'Шаблонные переменные {{name}} в тексте автоматически подставляются из хранилища платформы.')

    doc.save('ElemText.docx')
    print('✓ ElemText.docx')


def doc_elem_routes_navigator():
    doc = make_doc()
    add_title_block(doc, 'ElemRoutesNavigator', 'Виджет навигации по страницам приложения')

    add_h2(doc, 'Назначение')
    add_body(doc, (
        'ElemRoutesNavigator отображает список страниц/маршрутов платформы в виде '
        'горизонтального меню, вертикального меню, выпадающего списка или бургера. '
        'Поддерживает иерархию страниц, пагинацию, перетаскивание элементов '
        'в конструкторе и два режима открытия меню (клик / наведение).'
    ))

    add_h2(doc, 'Основные возможности')
    add_bullet(doc, '4 макета: horizontal / vertical / dropdown / burger', '')
    add_bullet(doc, 'Иерархия страниц с раскрываемыми уровнями и настраиваемым отступом', '')
    add_bullet(doc, 'Пагинация: прокрутка (scroll) или постраничная (pages)', '')
    add_bullet(doc, 'Режим открытия меню: click / hover', '')
    add_bullet(doc, 'Подсветка активной страницы', '')
    add_bullet(doc, 'Drag-to-reorder в редакторе (порядок страниц в `routesOrder`)', '')
    add_bullet(doc, 'Отключение отдельных страниц (`disabledPages`)', '')

    add_h2(doc, 'Параметры')

    add_h3(doc, 'Навигация')
    add_props_table(doc, [
        ('title',               'String',  'Навигация', 'Заголовок виджета'),
        ('showTitle',           'Boolean', 'true',      'Показывать заголовок'),
        ('orientation',         'String',  'horizontal','Макет: horizontal / vertical / dropdown / burger'),
        ('buttonAlignment',     'String',  'left',      'Выравнивание кнопок: left / center / right'),
        ('openMode',            'String',  'click',     'Открытие меню: click / hover'),
        ('highlightActivePage', 'Boolean', 'false',     'Подсвечивать текущую страницу'),
        ('showSlug',            'Boolean', 'false',     'Показывать слаги (ссылки) страниц'),
        ('dropdownText',        'String',  '\'Выберите страницу\'', 'Текст заглушки в выпадающем списке'),
        ('appJsonUrl',          'String',  '\'\'',      'URL файла app.json (если не задан — берётся автоматически)'),
        ('disabledPages',       'Array',   '[]',        'Слаги страниц, недоступных для перехода'),
        ('routesOrder',         'Array',   '[]',        'Порядок страниц (задаётся drag-and-drop в редакторе)'),
    ])

    add_h3(doc, 'Пагинация')
    add_props_table(doc, [
        ('enablePagination',      'Boolean', 'true',    'Включить пагинацию'),
        ('paginationType',        'String',  'scroll',  'Тип: scroll (непрерывная) / pages (постраничная)'),
        ('itemsPerPage',          'Number',  '6',       'Элементов на странице'),
        ('paginationActiveColor', 'String',  '#3b82f6', 'Цвет активного элемента пагинации'),
    ])

    add_h3(doc, 'Иерархия')
    add_props_table(doc, [
        ('enableHierarchy',  'Boolean', 'false', 'Включить иерархию страниц'),
        ('hierarchy',        'Object',  '{}',    'Объект иерархии {parentSlug: [childSlug, ...]}'),
        ('hierarchyIndent',  'Number',  '2.5',   'Отступ дочернего уровня (rem)'),
        ('navigateParents',  'Boolean', 'true',  'Переходить на страницу при клике на родительский элемент'),
    ])

    add_h3(doc, 'Стиль кнопок')
    add_props_table(doc, [
        ('buttonStyle',           'String',  'filled',   'Стиль: filled / outlined / ghost / underline'),
        ('activeColor',           'String',  '#3b82f6',  'Цвет активной страницы'),
        ('enableHoverColor',      'Boolean', 'true',     'Цвет при наведении'),
        ('hoverColor',            'String',  '#60a5fa',  'Цвет при наведении'),
        ('buttonBackgroundColor', 'String',  '#f3f4f6',  'Фон кнопки'),
        ('buttonPadding',         'Object',  '{0.75rem}','Отступ кнопки'),
        ('buttonGap',             'Object',  '{0.5rem}', 'Расстояние между кнопками'),
        ('buttonShadow',          'String',  '\'\'',     'Тень кнопки'),
    ])

    add_h3(doc, 'Контейнер')
    add_props_table(doc, [
        ('backgroundColor', 'String', '#ffffff',    'Фон контейнера'),
        ('textColor',       'String', '#1f2937',    'Цвет текста'),
        ('borderRadius',    'String', '0.375rem',   'Скругление'),
        ('boxShadow',       'String', '\'\'',       'Тень контейнера'),
    ])

    add_h3(doc, 'Типографика')
    add_props_table(doc, [
        ('fontSize',   'Object', '{0.875rem}', 'Размер шрифта {size, unit}'),
        ('fontFamily', 'String', 'inherit',    'Гарнитура шрифта'),
    ])

    add_h3(doc, 'Иконки')
    add_props_table(doc, [
        ('burgerIconOpen',      'String', '\'\'',    'MDI-иконка открытого бургер-меню'),
        ('burgerIconClosed',    'String', '\'\'',    'MDI-иконка закрытого бургер-меню'),
        ('burgerIconSize',      'String', '1.5rem',  'Размер иконки бургера'),
        ('expandIconExpanded',  'String', '\'\'',    'MDI-иконка раскрытого уровня'),
        ('expandIconCollapsed', 'String', '\'\'',    'MDI-иконка свёрнутого уровня'),
        ('expandIconSize',      'String', '0.75rem', 'Размер иконки уровня иерархии'),
    ])

    doc.save('ElemRoutesNavigator.docx')
    print('✓ ElemRoutesNavigator.docx')


def doc_elem_auto_logout():
    doc = make_doc()
    add_title_block(doc, 'ElemAutoLogout', 'Виджет автоматического выхода при бездействии пользователя')

    add_h2(doc, 'Назначение')
    add_body(doc, (
        'ElemAutoLogout отслеживает активность пользователя (движение мыши, клики, нажатия клавиш, прокрутка). '
        'При превышении таймаута бездействия показывает предупреждающий диалог с обратным отсчётом. '
        'Если пользователь не реагирует — выполняется выход из системы через AuthManager платформы. '
        'Опционально сохраняет URL текущей страницы в sessionStorage для возврата после входа.'
    ))

    add_h2(doc, 'Механизм работы')
    add_bullet(doc, 'Старт: отслеживание активности начинается при монтировании виджета', '')
    add_bullet(doc, 'Сброс: любое действие пользователя обнуляет таймер бездействия', '')
    add_bullet(doc, 'Предупреждение: за N секунд до выхода появляется диалог с обратным отсчётом', '')
    add_bullet(doc, 'Кнопка «Оставаться в системе»: скрывает диалог и сбрасывает таймер', '')
    add_bullet(doc, 'Выход: вызывается AuthManager.logout() платформы', '')
    add_bullet(doc, 'Сохранение URL: при preserveReturnUrl=true текущий адрес сохраняется в sessionStorage', '')
    add_note(doc, 'В режиме редактора (конструктора) трекинг не активируется — виджет показывает информационный плейсхолдер.')

    add_h2(doc, 'Параметры')

    add_h3(doc, 'Таймауты')
    add_props_table(doc, [
        ('timeoutSeconds', 'Number',  '1800', 'Время бездействия до выхода (по умолчанию 30 минут)'),
        ('timeoutUnit',    'String',  '\'\'', 'Единица измерения таймаута (s = секунды, m = минуты, h = часы). Если пусто — секунды'),
        ('warningEnabled', 'Boolean', 'true', 'Показывать предупреждение перед выходом'),
        ('warningDuration','Number',  '30',   'Время предупреждения до выхода'),
        ('warningUnit',    'String',  '\'\'', 'Единица измерения предупреждения (по умолчанию секунды)'),
    ])

    add_h3(doc, 'Тексты')
    add_props_table(doc, [
        ('labels.warningTitle',   'String', 'Сессия истекает',       'Заголовок диалога'),
        ('labels.warningMessage', 'String', 'Вы будете выведены через', 'Сообщение перед обратным отсчётом'),
        ('labels.stayLoggedIn',   'String', 'Оставаться в системе',   'Текст кнопки продления сессии'),
        ('labels.seconds',        'String', 'сек.',                   'Суффикс единицы времени в отсчёте'),
    ])

    add_h3(doc, 'Поведение')
    add_props_table(doc, [
        ('preserveReturnUrl', 'Boolean', 'true', 'Сохранять URL текущей страницы в sessionStorage для возврата после входа'),
    ])

    add_h3(doc, 'Оформление диалога')
    add_props_table(doc, [
        ('overlayColor',     'String', '#000000', 'Цвет затемняющей подложки'),
        ('overlayOpacity',   'Number', '45',      'Непрозрачность подложки (%)'),
        ('dialogBgColor',    'String', '\'\'',    'Фон диалогового окна'),
        ('dialogTextColor',  'String', '\'\'',    'Цвет текста диалога'),
        ('dialogRadius',     'Number', '8',       'Скругление окна (px)'),
        ('dialogMaxWidth',   'Number', '380',     'Максимальная ширина диалога (px)'),
        ('iconColor',        'String', '\'\'',    'Цвет иконки в диалоге'),
        ('fontFamily',       'String', '\'\'',    'Гарнитура шрифта диалога'),
    ])

    add_h3(doc, 'Оформление кнопки')
    add_props_table(doc, [
        ('btnBgColor',    'String', '\'\'', 'Фон кнопки «Оставаться»'),
        ('btnTextColor',  'String', '\'\'', 'Цвет текста кнопки'),
        ('btnRadius',     'Number', '8',    'Скругление кнопки (px)'),
        ('btnFontSize',   'String', '\'\'', 'Размер шрифта кнопки'),
        ('btnFontWeight', 'String', '\'\'', 'Толщина шрифта кнопки'),
        ('textTransform', 'String', 'none', 'Трансформация текста кнопки'),
        ('letterSpacing', 'String', '\'\'', 'Межбуквенный интервал кнопки'),
    ])

    add_h3(doc, 'Произвольный CSS')
    add_props_table(doc, [
        ('overlayCustomCss',  'String', '\'\'', 'CSS для подложки'),
        ('dialogCustomCss',   'String', '\'\'', 'CSS для диалогового окна'),
        ('dialogBtnCustomCss','String', '\'\'', 'CSS для кнопки в диалоге'),
    ])

    doc.save('ElemAutoLogout.docx')
    print('✓ ElemAutoLogout.docx')


def doc_elem_event_stack():
    doc = make_doc()
    add_title_block(doc, 'ElemEventStack', 'Виджет переключения состояний — контейнер со слотами-состояниями')

    add_h2(doc, 'Назначение')
    add_body(doc, (
        'ElemEventStack — контейнерный виджет, позволяющий показывать разный контент '
        'в зависимости от текущего состояния. Каждое состояние — именованный слот '
        'с произвольными дочерними виджетами и индивидуальным CSS. '
        'Переключение состояний происходит через события платформы. '
        'В редакторе все слоты видны одновременно (неактивные затемнены), '
        'в рантайме — только активный.'
    ))

    add_h2(doc, 'Концепция работы')
    add_bullet(doc, 'Определите список состояний в параметре `states` (массив {name, event})', '')
    add_bullet(doc, 'Поместите дочерние виджеты в каждое состояние в конструкторе', '')
    add_bullet(doc, 'Назначьте событие платформы для каждого состояния (поле event)', '')
    add_bullet(doc, 'При возникновении события — активируется соответствующий слот', '')
    add_bullet(doc, 'Начальное состояние задаётся через `activeState` (по умолчанию: «default»)', '')
    add_bullet(doc, 'Каждому состоянию можно назначить CSS через `stateStyles` (JSON)', '')

    add_h2(doc, 'Параметры')
    add_props_table(doc, [
        ('activeState',  'String', 'default',                      'Имя активного состояния при монтировании'),
        ('states',       'Array',  '[{name:\'default\',event:\'\'}]', 'Массив состояний: [{name: string, event: string}]'),
        ('stateStyles',  'String', '{}',                           'CSS для состояний (JSON): {\"имяСостояния\": \"css...\"}'),
        ('height',       'String', '100',                          'Высота контейнера'),
        ('heightUnit',   'String', '%',                            'Единица высоты: px / % / vh / и т.д.'),
    ])

    add_h2(doc, 'Пример конфигурации')
    add_body(doc, 'states:')
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(
        '[\n'
        '  { "name": "loading", "event": "data-loading" },\n'
        '  { "name": "loaded",  "event": "data-loaded"  },\n'
        '  { "name": "error",   "event": "data-error"   }\n'
        ']'
    )
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = DARK

    add_body(doc, '\nstateStyles:')
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.8)
    run2 = p2.add_run(
        '{\n'
        '  "loading": "opacity: 0.5; pointer-events: none;",\n'
        '  "error":   "border: 2px solid red;"\n'
        '}'
    )
    run2.font.name = 'Courier New'
    run2.font.size = Pt(9)
    run2.font.color.rgb = DARK

    add_note(doc, (
        'В режиме редактора все состояния видны одновременно. '
        'Неактивные слоты отображаются с пониженной прозрачностью (35%) '
        'и с именем состояния в виде бейджа для удобства редактирования.'
    ))

    doc.save('ElemEventStack.docx')
    print('✓ ElemEventStack.docx')


# ─── Run all ────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    import os
    os.chdir('/home/user/test4/2.1/docs')

    doc_elem_button()
    doc_elem_calendar()
    doc_elem_sticker()
    doc_elem_text()
    doc_elem_routes_navigator()
    doc_elem_auto_logout()
    doc_elem_event_stack()

    print('\nДокументация сохранена в /home/user/test4/2.1/docs/')
